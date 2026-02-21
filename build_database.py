from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from PIL import Image
import json, os, shutil

# Try to import pytesseract for OCR support
try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False
    print("⚠️  pytesseract not available - skipping image OCR. Install Tesseract-OCR for image support.")

print("📂 Loading college website data...")
all_docs = []

# Load scraped website data
with open("college_data.json", "r", encoding="utf-8") as f:
    raw_data = json.load(f)

for d in raw_data:
    all_docs.append(Document(
        page_content=d["content"],
        metadata={"source": d["source"]}
    ))
print(f"✅ Loaded {len(all_docs)} website pages")

# Load files from college_pdfs folder
pdf_folder = "./college_pdfs"
file_count = 0
if os.path.exists(pdf_folder):
    for filename in os.listdir(pdf_folder):
        filepath = os.path.join(pdf_folder, filename)
        
        # Handle PDF files
        if filename.lower().endswith(".pdf"):
            try:
                loader = PyPDFLoader(filepath)
                docs = loader.load()
                all_docs.extend(docs)
                file_count += len(docs)
                print(f"✅ Loaded PDF: {filename}")
            except Exception as e:
                print(f"❌ Failed to load {filename}: {e}")
        
        # Handle JPG/PNG image files - use OCR if available
        elif filename.lower().endswith((".jpg", ".jpeg", ".png")):
            if HAS_TESSERACT:
                try:
                    image = Image.open(filepath)
                    text = pytesseract.image_to_string(image)
                    if text.strip():
                        all_docs.append(Document(
                            page_content=text,
                            metadata={"source": f"Image: {filename}"}
                        ))
                        file_count += 1
                        print(f"✅ Loaded Image (OCR): {filename}")
                    else:
                        print(f"⚠️  No text found in image: {filename}")
                except Exception as e:
                    print(f"❌ Failed to read image {filename}: {e}")
            else:
                print(f"⚠️  Skipping image (OCR not available): {filename}")
        
        # Handle DOCX files
        elif filename.lower().endswith(".docx"):
            try:
                from docx import Document as DocxDocument
                docx = DocxDocument(filepath)
                text = "\n".join([para.text for para in docx.paragraphs])
                if text.strip():
                    all_docs.append(Document(
                        page_content=text,
                        metadata={"source": f"DOCX: {filename}"}
                    ))
                    file_count += 1
                    print(f"✅ Loaded DOCX: {filename}")
            except Exception as e:
                print(f"❌ Failed to load {filename}: {e}")

print(f"✅ Loaded {file_count} document pages from files")
print(f"📊 Total: {len(all_docs)} documents")

# Split into chunks
print("\n✂️  Splitting into chunks...")
splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=60)
chunks = splitter.split_documents(all_docs)
print(f"✅ Created {len(chunks)} chunks")

# Load embedding model
print("\n🔄 Loading embedding model (first time: 3-5 minutes)...")
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
print("✅ Embedding model ready!")

# Build database
print("\n💾 Building database...")
if os.path.exists("./college_db"):
    shutil.rmtree("./college_db")

vectordb = Chroma.from_documents(
    chunks,
    embeddings,
    persist_directory="./college_db"
)
vectordb.persist()
print("✅ Database built successfully!")
print("\n🎉 Done! Now run: streamlit run chatbot.py")
